import io
from types import SimpleNamespace

import docx
import pytest

from app.services import documents
from app.services.documents import UnsupportedDocumentError, extract_text


def test_extract_text_from_docx_joins_paragraphs():
    buffer = io.BytesIO()
    document = docx.Document()
    document.add_paragraph("Олена Ковальчук")
    document.add_paragraph("Бухгалтер, 8 років досвіду")
    document.save(buffer)

    text = extract_text(buffer.getvalue(), "resume.docx")

    assert text == "Олена Ковальчук\nБухгалтер, 8 років досвіду"


def test_extract_text_from_pdf_joins_pages(monkeypatch):
    fake_pages = [
        SimpleNamespace(extract_text=lambda: "Page one text"),
        SimpleNamespace(extract_text=lambda: "Page two text"),
    ]
    monkeypatch.setattr(documents, "PdfReader", lambda _stream: SimpleNamespace(pages=fake_pages))

    text = extract_text(b"irrelevant bytes", "resume.pdf")

    assert text == "Page one text\nPage two text"


def test_extract_text_rejects_unsupported_extension():
    with pytest.raises(UnsupportedDocumentError):
        extract_text(b"whatever", "resume.txt")
